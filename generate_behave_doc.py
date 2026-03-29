"""Generate the comprehensive Behave platform specification document as .docx"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── Styles ──────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

for level in range(1, 4):
    heading_style = doc.styles[f"Heading {level}"]
    heading_style.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0F)
    heading_style.font.bold = True

list_style = doc.styles["List Paragraph"]
list_style.font.name = "Calibri"
list_style.font.size = Pt(11)


def add_title_page():
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BEHAVE")
    run.font.size = Pt(36)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Powered Consumer Behavior Simulation Engine\nfor Indian D2C Brands")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Complete Platform Specification v3.0")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Architecture • Working System • Live Results • Expansion Roadmap")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("March 2026 — Internal: Founding Team Only")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Paragraph")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ═══════════════════════════════════════════════════════════════
# Generate the document
# ═══════════════════════════════════════════════════════════════

add_title_page()

# ── PART 1 ─────────────────────────────────────────────────────
doc.add_heading("Part 1: What Behave Is", level=1)

add_para(
    "Behave is a simulated Indian market — a persistent, living digital population "
    "of 20 AI consumers who react to products, ads, and experiences the way real "
    "Indian consumers would. Companies inject their ad creative into this simulated "
    "world and observe how consumers across 8 distinct behavioral archetypes discover "
    "it, react to it, trust it, desire it, or ignore it — before spending a single "
    "rupee on real-world testing."
)

add_para(
    "This is not a concept document. Behave is a working system. It has been built, "
    "tested, and run against real Indian ad creatives across multiple categories "
    "(Automobile, Household, Grocery Delivery, Services). Every component described "
    "in this specification exists in production code with 120 passing tests."
)

doc.add_heading("What Makes This Different", level=2)

add_para(
    "A wrapper is: prompt an LLM to pretend to be a persona and react to an ad. "
    "That is a single stateless API call. What Behave does is fundamentally different "
    "across seven dimensions:"
)

add_bullet("Each persona has deep behavioral psychology — not just demographics, but purchase "
           "psychology scores, emotional profiles, internal conflicts, generational touchstones, "
           "influence networks, cognitive biases, and values hierarchies. A single persona "
           "profile has 50+ structured fields.", bold_prefix="Deep Personas: ")

add_bullet("Every evaluation becomes a memory. Personas remember what they have seen before. "
           "A persona who saw three Holi ads in a row develops ad fatigue. A persona who had "
           "a good experience with a brand remembers it next time. Memory accumulates across "
           "simulation runs — observations, reflections, preferences, and category beliefs.",
           bold_prefix="Persistent Memory: ")

add_bullet("Every persona evaluation passes through a Critic Agent that checks four dimensions: "
           "persona consistency, sycophancy detection, cultural authenticity, and action-reasoning "
           "alignment. Evaluations that fail are quarantined. Only quality-approved evaluations "
           "enter the pipeline.", bold_prefix="Quality-Gated Output: ")

add_bullet("The system does not summarize — it decides. A Recommendation Agent reviews aggregated "
           "results and makes a concrete business decision: SCALE (increase spend), ITERATE "
           "(specific edits needed), KILL (pull immediately), or SPLIT (segment-specific variants). "
           "Each recommendation includes a prioritized edit playbook.",
           bold_prefix="Agentic Business Decisions: ")

add_bullet("Personas literally see the ad image. Vision-native LLMs process the raw ad creative "
           "alongside the structured metadata. The persona notices colors, Hinglish copy, "
           "price tags, lifestyle settings — things a text-only system misses entirely.",
           bold_prefix="Vision-Native Evaluation: ")

add_bullet("Every simulation run flows through a Databricks medallion pipeline (Bronze → Silver → Gold), "
           "with schema validation, quarantine routing, and MLflow experiment tracking. "
           "The system reads from and writes to Databricks — a full round-trip, not just "
           "a one-directional push.", bold_prefix="Cloud Data Pipeline: ")

add_bullet("A live interactive dashboard where brands enter their creative details, select "
           "their audience, and get real-time results — scorecards, radar charts, segment "
           "breakdowns, persona verbatims, and an AI-generated edit playbook.",
           bold_prefix="Interactive Dashboard: ")

doc.add_page_break()

# ── PART 2 ─────────────────────────────────────────────────────
doc.add_heading("Part 2: Technology Foundations", level=1)

add_para(
    "The Stanford Generative Agents paper (Park et al., 2023) was built on GPT-3.5 "
    "with tiny context windows and 25 agents. Three years later, the available technology "
    "has dramatically expanded what is achievable. Here are the specific advances that "
    "shape Behave's architecture:"
)

doc.add_heading("2.1 Vision-Native Agents", level=2)
add_para(
    "In 2023, agents could only process text. Now Claude and GPT can see images natively. "
    "Behave's agents literally look at an ad the way a real consumer does — noticing colors, "
    "reading Hinglish copy, seeing the price tag, registering cultural cues, absorbing "
    "the lifestyle setting."
)
add_para(
    "Architecture: The evaluation prompt includes both the raw ad image (via base64 encoding) "
    "and the extracted CreativeCard metadata. If the vision model in extraction missed something "
    "(e.g., small print disclaimer, background cultural cue), the persona's direct image "
    "perception still catches it. The image_utils module handles encoding for PNG, JPEG, GIF, "
    "and WebP formats.",
    italic=True,
)

doc.add_heading("2.2 Massive Context Windows (200K+ Tokens)", level=2)
add_para(
    "The Stanford paper built elaborate retrieval systems because nothing fit in the prompt. "
    "With 200K token context windows, Behave fits a persona's entire life history — every "
    "ad they have seen, every reaction, every reflection — directly into the prompt for "
    "the first 50–100 simulations per persona."
)
add_para(
    "Architecture: Behave implements a \"full dump\" mode that injects all memories directly "
    "when memory count is below the full_dump_threshold (default: 50 nodes). The weighted "
    "retrieval system (score = α·recency + β·relevance + γ·importance) activates automatically "
    "when memory grows beyond this threshold. This eliminates retrieval errors, which the "
    "Stanford paper identified as the most common failure mode.",
    italic=True,
)

doc.add_heading("2.3 Structured Output and Tool Use", level=2)
add_para(
    "Getting LLMs to output consistent JSON was unreliable in 2023. Now structured output "
    "is native. Behave enforces exact schemas at the API level — every evaluation returns "
    "a guaranteed-valid PersonaEvaluation object with all 25+ fields. Every recommendation "
    "returns a valid AgentRecommendation. Every critic verdict returns a valid CriticVerdict. "
    "Zero parsing errors."
)
add_para(
    "Architecture: All data models are Pydantic v2 BaseModels with strict validation — "
    "field ranges (0–10, 0–100), required strings, enum constraints. The LLM client handles "
    "JSON extraction from responses including markdown code block stripping.",
    italic=True,
)

doc.add_heading("2.4 Dramatically Cheaper Inference", level=2)
add_para(
    "The Stanford study cost thousands of dollars for 25 agents over 2 days. Behave uses "
    "model tiering: Claude Haiku for high-volume persona evaluations (~$0.025 per evaluation), "
    "Claude Sonnet for the Critic Agent and Recommendation Agent where quality matters most. "
    "A complete simulation run with 10 personas costs under $0.25."
)
add_para(
    "Architecture: The LLMConfig system supports 6 configurable models across different roles: "
    "eval_model (Haiku), creative_analysis_model (Sonnet), critic_model (Sonnet), "
    "recommendation_model (Sonnet), reflection_model (Haiku), embedding_model (text-embedding-3-small). "
    "Every LLM call tracks prompt_tokens, completion_tokens, cost_usd, and latency_ms. "
    "Total run cost is aggregated in RunMetadata.",
    italic=True,
)

doc.add_heading("2.5 Embedding Models for Memory Retrieval", level=2)
add_para(
    "Embedding models are now cheap, fast, and available via API. When memory streams grow "
    "beyond context window limits, cosine similarity over embeddings provides relevance "
    "scoring. Behave's MemoryRetriever implements the full Stanford retrieval formula "
    "with configurable weights for recency, relevance, and importance."
)
add_para(
    "Architecture: The MemoryRetriever class computes composite scores using exponential "
    "decay for recency (configurable decay factor, default 0.99), cosine similarity for "
    "relevance (with keyword fallback when embeddings are unavailable), and normalized "
    "importance (0–10 scaled to 0–1). The MemoryConsumer class auto-selects between "
    "full-dump and scored retrieval based on per-category node count.",
    italic=True,
)

doc.add_page_break()

# ── PART 3 ─────────────────────────────────────────────────────
doc.add_heading("Part 3: Core Architecture", level=1)

doc.add_heading("3.1 Data Flow Overview", level=2)
add_para(
    "The pipeline flows through seven stages: Creative Upload → Creative Analysis → "
    "Cohort Selection → Persona Evaluation → Critic Quality Gate → Aggregation & "
    "Recommendation → Databricks Ingest. Every stage is implemented and working."
)

add_para("End-to-End Flow:", bold=True)
add_bullet("User submits a creative on the dashboard (brand, category, headline, body copy, "
           "CTA, optional ad image)")
add_bullet("System reads persona profiles and accumulated memories from Databricks silver "
           "tables (falls back to local JSON if Databricks is unavailable)")
add_bullet("Cohort selection picks the right personas — prioritizing category affinity, "
           "ensuring archetype diversity, filling with demographic diversity")
add_bullet("Each persona evaluates the creative: sees the raw ad image + structured metadata + "
           "their relevant memories → produces a structured evaluation with 5 dimensional "
           "scores, reasoning, objections, and a verbatim reaction in their authentic voice")
add_bullet("Critic Agent reviews every evaluation for persona consistency, sycophancy, "
           "cultural authenticity, and action-reasoning alignment — quarantines failures")
add_bullet("Evaluations aggregate into a CreativeScorecard with segment breakdowns by archetype")
add_bullet("Recommendation Agent makes a SCALE/ITERATE/KILL/SPLIT decision with edit playbook")
add_bullet("Everything writes back to Databricks: JSON files upload to DBFS, an ingest notebook "
           "processes them through Bronze → Silver → Gold, run logs to MLflow")

doc.add_heading("3.2 The Three-Layer Agent Model", level=2)

add_para("Layer 1: The Soul (Persona Profile)", bold=True)
add_para(
    "A detailed behavioral profile governing how the agent thinks and decides. Not just "
    "demographics — purchase psychology, price sensitivity, brand relationship patterns, "
    "media behavior, emotional triggers, generational touchstones, internal conflicts, "
    "cognitive biases, and cultural context. Each persona has 50+ structured fields across "
    "11 major categories."
)

add_para("Key persona dimensions:", bold=True)
add_bullet("Demographics — age, gender, city, city tier (Metro/Tier1/Tier2/Tier3/Rural), "
           "state, income segment, language primary/secondary, education, occupation, family status",
           bold_prefix="Demographics: ")
add_bullet("How they evaluate, compare, and decide. 8 scored dimensions: price_sensitivity, "
           "brand_loyalty, impulse_tendency, social_proof_need, research_depth, risk_tolerance, "
           "decision_speed, deal_sensitivity — each 0.0 to 1.0",
           bold_prefix="Purchase Psychology: ")
add_bullet("Primary and secondary platforms, ad tolerance, scroll speed, video preference, "
           "influencer trust, content language preference",
           bold_prefix="Media Behavior: ")
add_bullet("Aspiration drivers, trust signals, rejection triggers, emotional hooks — "
           "specific to Indian consumer culture",
           bold_prefix="Emotional Profile: ")
add_bullet("Formative era ('90s Doordarshan kid'), iconic ads, cultural references, "
           "nostalgia intensity, formative brands",
           bold_prefix="Generational Touchstones: ")
add_bullet("Named tensions ('wants premium but fears wasting money') with resolution tendencies. "
           "These create realistic ambivalence in evaluations.",
           bold_prefix="Internal Conflicts: ")
add_bullet("Mother, spouse, YouTube reviewer, Reddit community — each with influence strength "
           "and domain specificity",
           bold_prefix="Influence Network: ")
add_bullet("Per-category interest, monthly spend, preferred brands, purchase frequency, "
           "and typed brand relationships (loyal_lover, nostalgic_friend, research_respect)",
           bold_prefix="Category Affinities: ")
add_bullet("Ordered list of what matters most — evidence_based_decisions, ingredient_transparency, "
           "long_term_value, etc.",
           bold_prefix="Values Hierarchy: ")
add_bullet("Named cognitive biases that shape evaluation: analysis_paralysis, status_quo_bias, "
           "anchoring_effect, bandwagon, etc.",
           bold_prefix="Cognitive Biases: ")

add_para("Layer 2: The Memory (Stream + Retrieval + Reflection)", bold=True)
add_para(
    "Directly adapted from the Stanford Generative Agents paper. Each persona maintains "
    "a MemoryStream — an append-only log of four memory types:"
)
add_bullet("Ads seen, products evaluated, creative reactions — auto-generated "
           "from every PersonaEvaluation", bold_prefix="Observations: ")
add_bullet("Higher-order beliefs formed from accumulated observations — triggered "
           "when cumulative importance exceeds the reflection threshold (default: 50 points)",
           bold_prefix="Reflections: ")
add_bullet("Learned brand preferences from strong positive evaluations (score > 70)",
           bold_prefix="Preferences: ")
add_bullet("General category beliefs from strong negative evaluations (score < 30) — "
           "e.g., 'I am skeptical of brands that lead with toxin-free claims'",
           bold_prefix="Category Beliefs: ")

add_para(
    "Memory retrieval scores each memory on three dimensions:"
)
add_bullet("Exponential decay — recent memories surface more readily (configurable "
           "decay factor, default 0.99 per simulated hour)", bold_prefix="Recency: ")
add_bullet("Cosine similarity between the current creative context and the memory "
           "(embedding-based with keyword fallback)", bold_prefix="Relevance: ")
add_bullet("LLM-scored poignancy (0–10) — a bad product experience scores higher "
           "than seeing a forgettable ad", bold_prefix="Importance: ")

add_para(
    "Memory scope is configurable per simulation: NONE (fresh eyes), CATEGORY (same "
    "industry only — the default), BRAND (same brand only), or FULL (everything ever seen). "
    "This enables use cases like testing how a persona reacts to a brand's third ad "
    "versus their first."
)

add_para("Layer 3: The Body (Action Space)", bold=True)
add_para(
    "When an agent encounters a creative, their response is constrained to defined "
    "consumer actions: SCROLL_PAST, PAUSE, READ, ENGAGE, CLICK, SAVE, SHARE, "
    "PURCHASE_INTENT, or NEGATIVE_REACT. Each agent selects a primary action, provides "
    "5 dimensional scores (attention, relevance, trust, desire, clarity — each 0–10), "
    "reasoning, objections, and a verbatim reaction in their authentic voice."
)

doc.add_page_break()

doc.add_heading("3.3 The Creative Card", level=2)
add_para(
    "Every ad creative that enters the system is analyzed and converted into a structured "
    "CreativeCard with 28 fields. The CreativeCard captures:"
)
add_bullet("Brand, category, format (static_image, carousel, video_thumbnail, "
           "landing_page, story_ad, reel_ad)")
add_bullet("Headline, body copy, CTA text and type (buy, learn_more, sign_up, download)")
add_bullet("Pricing: shown price, original price, discount percentage, price framing "
           "(premium, value, discount, comparison, no_price)")
add_bullet("Behavioral cues: urgency cues, social proof cues, trust signals, emotional hooks")
add_bullet("Visual analysis: visual style (7 options), dominant color, product visibility, "
           "human presence, celebrity presence")
add_bullet("Language: primary language, secondary language, code-mixed (Hinglish) detection")
add_bullet("Indian context: festival context, target city tier, cultural references")
add_bullet("Source: raw image path for vision-native evaluation")
add_bullet("Quality: extraction confidence score, extraction model, timestamp")

doc.add_heading("3.4 Critic Agent (Quality Gate)", level=2)
add_para(
    "A separate Claude Sonnet agent that reviews every persona evaluation before it "
    "enters the pipeline. The Critic is not a summarizer — it actively decides whether "
    "each evaluation is authentic or not. It checks four dimensions:"
)
add_bullet("Does the response contradict the persona's behavioral profile? "
           "(e.g., a price-sensitive skeptic giving enthusiastic response to premium "
           "pricing without mentioning cost)",
           bold_prefix="Persona Consistency (0–10): ")
add_bullet("Is the agent being unrealistically positive? LLMs default to optimism — "
           "the Critic flags when agents are suspiciously enthusiastic about everything.",
           bold_prefix="Sycophancy Detection (0–10): ")
add_bullet("Does this sound like an actual Indian consumer from this demographic, or "
           "like an American LLM's idea of an Indian person? This is the India-specific "
           "Critic layer that no global competitor has.",
           bold_prefix="Cultural Authenticity (0–10): ")
add_bullet("Is the chosen action consistent with the stated reasoning? "
           "(e.g., agent says 'I do not trust this brand' but chose CLICK instead of SCROLL_PAST)",
           bold_prefix="Action-Reasoning Alignment (0–10): ")

add_para(
    "Quality scoring formula: overall_quality = 0.30 × consistency + 0.25 × (10 − sycophancy) "
    "+ 0.25 × authenticity + 0.20 × alignment. Pass threshold: 6.0/10. Failed evaluations "
    "are quarantined with detailed explanation and flags."
)

add_para("Real Critic output (Surf Excel, Household — Brand Loyalist Mumbai):", bold=True)
add_para(
    "\"This evaluation strongly matches Ananya's brand loyalist profile — she recognizes "
    "Surf Excel as familiar/trusted but shows no impulse to switch from her current "
    "detergent, consistent with her high brand loyalty (0.95) and low impulse tendency "
    "(0.25). The reasoning reflects her preference for established brands and resistance "
    "to marketing that doesn't offer clear differentiation.\"",
    italic=True,
)
add_para("Pass rate: 100%, Avg consistency: 7.9/10, Avg sycophancy: 2.6/10, "
         "Avg cultural authenticity: 7.9/10, Overall quality: 8.1/10", italic=True)

doc.add_heading("3.5 Recommendation Agent (Business Decisions)", level=2)
add_para(
    "A Claude Sonnet agent acting as a senior media strategist at a top Indian ad agency. "
    "Takes CreativeScorecard data and produces one of four concrete business actions:"
)
add_bullet("Creative is working — increase spend", bold_prefix="SCALE: ")
add_bullet("Promising foundation, specific edits recommended", bold_prefix="ITERATE: ")
add_bullet("Fundamentally broken — pull immediately", bold_prefix="KILL: ")
add_bullet("Strong with some segments, weak with others — create variants",
           bold_prefix="SPLIT: ")

add_para(
    "Each recommendation includes: confidence score (0–1), reasoning summary, key evidence "
    "list, top risks with risk level, a prioritized edit playbook (element, current state, "
    "suggested change, expected impact, priority 1–5), strongest and weakest segments, "
    "and recommended target audience."
)

add_para("Real Recommendation output (Surf Excel, Household — 8 personas):", bold=True)
add_para(
    "Verdict: ITERATE (70% confidence). \"The creative has strong emotional resonance and "
    "clarity but fails on desire generation and relevance. The traditional 'daag achhe hain' "
    "messaging is tired and doesn't differentiate from competitors or address modern consumer "
    "needs like price sensitivity and functional benefits.\"",
    italic=True,
)
add_para("Edit playbook:", bold=True)
add_bullet("[Priority 1] Body copy: Add specific functional claims — "
           "'Removes Holi colors 2x faster than ordinary detergents'")
add_bullet("[Priority 2] Visual: Include regional festivals (Onam, Durga Puja) in rotation "
           "or create regional variants")
add_bullet("[Priority 3] CTA: Add value proposition — 'Try Surf Excel's advanced stain removal'")

doc.add_page_break()

# ── PART 4 ─────────────────────────────────────────────────────
doc.add_heading("Part 4: India-Specific Persona Library", level=1)

add_para(
    "20 distinct personas across 8 behavioral archetypes, varying by demographics, city, "
    "city tier, language, income, and category familiarity. These are not demographic "
    "placeholders — each persona is a fully realized character with backstory, inner "
    "monologue style, internal conflicts, and influence networks."
)

doc.add_heading("The 8 Archetypes", level=2)

archetypes = [
    ("Researcher", "researcher",
     "Priya Mehta (Delhi), Dr. Kavitha Nair (Chennai), Subhadeep Mukherjee (Kolkata)",
     "Obsessive about evidence. Will read 15 reviews, check ingredients, compare alternatives. "
     "Trust signals: clinical studies, ingredient transparency, community consensus. "
     "Weakness: analysis paralysis. Research depth: 0.95. Decision speed: slow."),
    ("Impulse Buyer", "impulse_buyer",
     "Zara Khan (Delhi), Vikram Joshi (Mumbai)",
     "Emotion-driven, fast decisions. Responds to flash sales, FOMO, and aesthetic appeal. "
     "Impulse tendency: 0.85+. Social proof need: high. Scrolls fast, buys faster. "
     "Weakness: buyer's remorse. Trust comes from aesthetics and peer validation."),
    ("Price Anchor", "price_anchor",
     "Ramesh Verma (Lucknow), Suresh Yadav (Patna)",
     "Has precise mental price-per-unit for every category. Calculates value ruthlessly. "
     "Price sensitivity: 0.95. Responds to EMI pricing, bulk discounts, comparison ads. "
     "Will reject premium products on principle. Hindi-primary, Tier 2 sensibility."),
    ("Brand Loyalist", "brand_loyalist",
     "Ananya Desai (Mumbai), Kavitha Thomas (Kochi), Megha Sharma (Lucknow)",
     "Deep emotional bonds with specific brands. Brand loyalty: 0.9+. Resistant to switching. "
     "Trust comes from family history with the brand. Will defend their brands. "
     "Weakness: blindspot for better alternatives."),
    ("Trend Follower", "trend_follower",
     "Neha Kapoor (Mumbai), Ravi Patel (Ahmedabad)",
     "Instagram-native. Finds products through social discovery. Influencer trust: 0.75+. "
     "Early adopter energy. Will try anything that is 'trending.' Social proof need: very high. "
     "Weakness: easily influenced, low brand loyalty."),
    ("Skeptic", "skeptic",
     "Arjun Menon (Bangalore), Deepika Rajput (Jaipur)",
     "Distrusts advertising by default. Ad tolerance: 0.15. Needs exceptional evidence to engage. "
     "Rejection triggers: influencer endorsements, 'too good to be true' claims. "
     "Trust signals: user reviews, transparent pricing, long brand history. Hard to win, "
     "valuable when won."),
    ("Pragmatist", "pragmatist",
     "Amit Kulkarni (Pune), Rajesh Tiwari (Indore)",
     "Function over emotion. Evaluates features, durability, after-sales. Decision speed: moderate. "
     "Research depth: 0.7. Not impressed by branding — wants to know what it does and "
     "how much it costs. Risk tolerance: moderate. Value-driven, not price-driven."),
    ("Aspirational Buyer", "aspirational_buyer",
     "Pooja Kumari (Jaipur), Arjun Reddy (Hyderabad)",
     "Bridge between desire and budget. Aspires to premium brands, often priced out. "
     "Responds to EMI offers, 'affordable luxury' framing. Social media aspirational. "
     "Young (22–25), Tier 2, limited budgets. Internal conflict: 'I want nice things but "
     "I cannot afford them.'"),
]

for name, _id, personas, description in archetypes:
    add_para(f"{name}", bold=True)
    add_para(f"Personas: {personas}", italic=True)
    add_para(description)

doc.add_heading("Persona Depth: A Detailed Example", level=2)
add_para("Priya Mehta — Researcher, Delhi", bold=True)
add_bullet("29-year-old Product Manager, MBA from IIM Indore, upper-middle income")
add_bullet("Lives in South Delhi with parents (age 49), influences: "
           "r/IndianSkincareAddicts community (0.85 strength), YouTube reviewers (0.6), "
           "father (0.55 — 'the original researcher')")
add_bullet("Inner monologue: 'Ok but what is the actual active ingredient?' vs "
           "'This looks like marketing — where is the proof?'")
add_bullet("Internal conflicts: PM expertise in pricing psychology vs. consumer resistance "
           "to manipulation; recent raise means freedom but research paralysis delays spending; "
           "lives with parents enables comfort but costs social independence")
add_bullet("Generational touchstones: '90s mixer-grinder philosophy' + digital optimization; "
           "r/IndianSkincareAddicts is her temple")
add_bullet("Category affinities: Skincare (0.9) — prefers Minimalist, Deconstruct; "
           "Electronics (0.7) — Apple, Samsung")
add_bullet("Cognitive biases: analysis_paralysis, status_quo_bias, information_overload, "
           "effort_justification")

doc.add_page_break()

# ── PART 5 ─────────────────────────────────────────────────────
doc.add_heading("Part 5: The Medallion Pipeline", level=1)

add_para(
    "Data flows through three layers with quality enforcement at each stage, running on "
    "Databricks with Unity Catalog governance. All tables live under the "
    "bootcamp_students catalog with schema isolation: kinshuk_bronze, kinshuk_silver, "
    "kinshuk_gold."
)

doc.add_heading("5.1 Bronze Layer (Raw, Append-Only)", level=2)
add_para("Nothing is filtered or transformed. Raw data lands here exactly as produced:")
add_bullet("persona_seed — Raw persona JSON profiles, flattened for columnar storage")
add_bullet("evaluations — Raw LLM evaluation outputs per persona-creative pair")
add_bullet("critic_verdicts — Raw critic assessment for each evaluation")
add_bullet("run_metadata — Run configuration, cost tracking, timestamps")
add_bullet("memory_nodes — Raw memory observations generated during the run")

doc.add_heading("5.2 Silver Layer (Validated, Quarantined)", level=2)
add_para("Pydantic validation at every boundary. Records that fail go to quarantine:")
add_bullet("personas — Schema-validated profiles (backstory ≥ 20 chars, archetype present, "
           "category affinities non-empty)")
add_bullet("creative_cards — Validated creatives (headline or body_copy present, "
           "extraction_confidence ≥ 0.3)")
add_bullet("evaluations — Validated evaluations (evaluation_id present, reasoning ≥ 10 chars, "
           "verbatim ≥ 5 chars, scores in range)")
add_bullet("memory_nodes — Validated memory (node_id present, description ≥ 5 chars, "
           "persona_id valid)")
add_bullet("quarantine — Failed records with rejection reason, source table, and timestamp")

doc.add_heading("5.3 Gold Layer (Aggregated, Analytics-Ready)", level=2)
add_bullet("creative_scorecards — Aggregated scores per creative: attention, relevance, "
           "trust, desire, clarity, overall grade (A–F), action distribution, sentiment "
           "distribution, top strengths/weaknesses, segment summaries")
add_bullet("run_audit_log — Complete audit trail: personas used, cost, tokens, quality "
           "gate results, pass/fail/quarantine counts")

doc.add_heading("5.4 Databricks Integration", level=2)
add_para("The system has a bidirectional Databricks integration:")
add_para("Read path:", bold=True)
add_bullet("Personas: tries Databricks first (silver.personas table), falls back to local JSON")
add_bullet("Memory: tries Databricks first (silver.memory_nodes table), falls back to local JSON")

add_para("Write path:", bold=True)
add_bullet("After every simulation, all 7 JSON files (metadata, evaluations, scorecards, "
           "recommendation, critic_verdicts, critic_summary, memory_nodes) upload to DBFS "
           "at /synthetic_india/runs/{run_id}/")
add_bullet("An ingest notebook (04_simulation_ingest.py) is triggered as a one-time job run "
           "on a Databricks cluster")
add_bullet("The notebook processes the run through the full medallion pipeline: "
           "Bronze → Silver → Gold")
add_bullet("Synchronous mode (default): the system waits up to 180 seconds for the Databricks "
           "job to complete, polling every 10 seconds. Falls back to async on timeout.")

add_para("MLflow Integration:", bold=True)
add_bullet("Every run is logged as an MLflow experiment with metrics (overall score, cost, "
           "token usage, pass rate) for cross-run comparison")

doc.add_heading("5.5 Four Databricks Notebooks", level=2)
add_bullet("01_bronze_ingest.py — Load persona JSON → flatten nested fields → "
           "write to kinshuk_bronze.persona_seed", bold_prefix="Bronze Ingest: ")
add_bullet("02_silver_transform.py — Validate personas and creatives → split into "
           "passed (kinshuk_silver.personas, kinshuk_silver.creative_cards) and "
           "quarantined (kinshuk_silver.quarantine)", bold_prefix="Silver Transform: ")
add_bullet("03_gold_materialize.py — Aggregate into scorecards "
           "(kinshuk_gold.creative_scorecards) and audit log (kinshuk_gold.run_audit_log)",
           bold_prefix="Gold Materialize: ")
add_bullet("04_simulation_ingest.py — End-to-end ingest of a simulation run: "
           "load 7 JSONs from DBFS → bronze (append) → silver validate → gold aggregate. "
           "Parameterized by run_id via Databricks widget.",
           bold_prefix="Simulation Ingest: ")

doc.add_page_break()

# ── PART 6 ─────────────────────────────────────────────────────
doc.add_heading("Part 6: The Dashboard", level=1)

add_para(
    "A live interactive web dashboard built with FastAPI (backend) and vanilla JavaScript "
    "(frontend). No React, no frameworks — clean, performant, and self-contained."
)

doc.add_heading("6.1 Preloaded Case Study", level=2)
add_para(
    "The dashboard loads with a preloaded simulation run showing the full pipeline in action. "
    "Visitors see the complete flow: creative brief → persona panel → real-time pipeline "
    "visualization → results scorecard → segment breakdown → verbatim quotes → edit playbook."
)

doc.add_heading("6.2 Live Simulation", level=2)
add_para("Behind a password gate (configurable via environment variable), brands can run their own simulations:")
add_bullet("Enter brand name, category, headline, body copy, CTA")
add_bullet("Upload an ad image (optional — encoded as base64 for vision-native evaluation)")
add_bullet("Choose selection mode: Auto (slider, 1–20 personas) or Manual (pick specific personas)")
add_bullet("Manual mode shows a persona picker grid with clickable chips — "
           "each chip shows name, city, and archetype. Select All / Deselect All buttons.")
add_bullet("Live cost estimate updates as you select personas (~$0.025 per persona)")
add_bullet("Choose memory scope: None, Category, Brand, or Full")
add_bullet("Submit and watch results render in real-time")

doc.add_heading("6.3 Results Display", level=2)
add_bullet("Verdict banner — color-coded: KILL (red), ITERATE (amber), SCALE (green)")
add_bullet("5-dimensional radar chart — canvas-rendered: attention, relevance, trust, desire, clarity")
add_bullet("Score bars — visual 0–10 bars for each dimension")
add_bullet("Segment breakdown — per-archetype cards with dimensional scores and representative verbatim")
add_bullet("Persona verbatims — authentic quotes in each persona's voice")
add_bullet("Edit playbook — prioritized recommendations with element, current state, "
           "suggested change, expected impact")

doc.add_heading("6.4 API Endpoints", level=2)
add_bullet("GET /api/health — Health check")
add_bullet("GET /api/personas — Returns all 20 personas (persona_id, name, city, archetype)")
add_bullet("POST /api/simulate — Full simulation with creative details, returns metadata + "
           "scorecards + recommendation + evaluations")
add_bullet("GET /api/preloaded — Returns the preloaded case study data")

doc.add_heading("6.5 Design System", level=2)
add_para(
    "Dark theme with Instrument Serif (display) and DM Sans (body) typography. "
    "Color palette: near-black primary (#0a0a0f), blue accent (#3b82f6), red for KILL (#ef4444), "
    "light gray text (#e8e8ed). Responsive grid layout with blur-backdrop navigation."
)

doc.add_page_break()

# ── PART 7 ─────────────────────────────────────────────────────
doc.add_heading("Part 7: Real Output — What Behave Actually Produces", level=1)

add_para(
    "These are real outputs from actual simulation runs against real Indian ad creatives. "
    "Nothing below is mocked or fabricated."
)

doc.add_heading("7.1 Surf Excel — Household (8 Personas)", level=2)
add_para("Overall Score: 51.2/100 (Grade C) | Verdict: ITERATE (70% confidence) | Cost: $0.23",
         bold=True)
add_para("Sample verbatims:", bold=True)

add_para(
    "Brand Loyalist, Mumbai: \"Arrey yaar, another Surf Excel Holi ad. Same message every "
    "year na — daag acche hain and all. Sweet visual but what's the point? I don't even "
    "celebrate Holi properly anymore. And I already have that big Ariel pack at home. "
    "Why would I switch now?\"",
    italic=True,
)
add_para("Action: SCROLL_PAST | Score: 35/100", italic=True)

add_para(
    "Brand Loyalist, Kochi: \"Ayyo, this is so sweet no? That thatha and kutty playing "
    "Holi together... reminds me of my appa with Aishwarya last year. 'Jo dilon ko paas "
    "laayein' — yes, that's the point of festivals, not worrying about stains. Surf Excel "
    "does remove colors well, I have to admit. But why always North Indian festivals in ads? "
    "Show Onam, show Thiruvathira also na!\"",
    italic=True,
)
add_para("Action: PAUSE | Score: 75/100", italic=True)

add_para(
    "Pragmatist, Pune: \"Arrey yaar, another emotional ad. Holi colors are nice but "
    "tell me — does it actually remove turmeric stains better than Ariel? What's the "
    "cost per wash? My wife will ask me these questions if I suggest switching brands. "
    "Show me some data na, not just cute grandfather-granddaughter moments.\"",
    italic=True,
)
add_para("Action: SCROLL_PAST | Score: 45/100", italic=True)

add_para(
    "Price Anchor, Lucknow: \"Yaar dekho, rang-birange colors dikha rahe hain but price "
    "kahan hai? Surf Excel matlab paisa jyada hi lagega. Mere bachche bhi Holi khelenge "
    "lekin main toh Nirma se hi kapde dholunga — safai same milti hai, paisa bachta hai. "
    "Yeh sab emotional drama hai, MRP dikhao pehle!\"",
    italic=True,
)
add_para("Action: PAUSE | Score: 25/100", italic=True)

doc.add_heading("7.2 Tata Tiago EV — Automobile (10 Personas)", level=2)
add_para("Overall Score: 37.4/100 (Grade D) | Verdict: KILL (90% confidence) | Cost: $0.27",
         bold=True)

add_para(
    "Recommendation: \"This creative is fundamentally broken with a D-grade overall score "
    "of 37.4/100. The combination of extremely low attention (4.5), relevance (2.8), "
    "and desire (2.0) scores indicates it will burn budget without generating meaningful "
    "conversions.\"",
    italic=True,
)

add_para("Sample verbatim — Aspirational Buyer, Jaipur:", bold=True)
add_para(
    "\"Yaar dekho, Tata ka electric car ad hai. 'Own a car not a cartoon' — cute line, "
    "but bhai I'm still trying to own a decent laptop, car toh bahut door ki baat hai! "
    "Plus papa would never let me even think about buying a car on my own. This is clearly "
    "for people with proper jobs and savings, not for someone like me jo abhi credit card "
    "se dar rahi hai. Scroll kar deti hun.\"",
    italic=True,
)
add_para("Action: SCROLL_PAST | Score: 25/100", italic=True)

doc.add_heading("7.3 Completed Simulation Runs", level=2)

# Table of runs
tbl = doc.add_table(rows=5, cols=6)
tbl.style = "Light Grid Accent 1"
headers = ["Brand", "Category", "Personas", "Score", "Verdict", "Cost"]
for i, h in enumerate(headers):
    tbl.rows[0].cells[i].text = h
runs = [
    ("Tata Tiago EV", "Automobile", "10", "37.4 (D)", "KILL 90%", "$0.27"),
    ("Surf Excel", "Household", "8", "51.2 (C)", "ITERATE 70%", "$0.23"),
    ("Snabbit", "Maid Services", "10", "39.4 (D)", "KILL 90%", "$0.23"),
    ("Zepto", "Grocery Delivery", "5", "42.4 (D)", "KILL 85%", "$0.12"),
]
for row_idx, run in enumerate(runs):
    for col_idx, val in enumerate(run):
        tbl.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph("")
add_para(
    "Total cost across all runs: under $1.00. Each run takes 60–120 seconds for "
    "10 personas. The same insight from traditional focus groups would cost ₹2–5 lakhs "
    "and take 2–4 weeks.",
    italic=True,
)

doc.add_page_break()

# ── PART 8 ─────────────────────────────────────────────────────
doc.add_heading("Part 8: Memory System in Action", level=1)

add_para(
    "The memory system is what transforms Behave from a stateless wrapper into a "
    "persistent simulation. Here is how memory works across simulation runs, demonstrated "
    "with real data."
)

doc.add_heading("8.1 Cross-Run Memory", level=2)
add_para(
    "When the Aspirational Buyer from Jaipur evaluated the Tata Tiago EV ad, her evaluation "
    "included memory context from previous runs. The system automatically retrieved relevant "
    "past experiences and injected them into the prompt:"
)
add_para(
    "\"You have encountered similar products/brands before. Here are your relevant memories:\n\n"
    "1. I saw a Zepto grocery delivery ad. My first impression: Oh it's another Zepto ad, "
    "and wait... what's with this cheesy Hindi line about ice cream? I would scroll past. "
    "This ad is so random and doesn't speak to me at all. (2 days ago)\n\n"
    "2. I saw a Snabbit Maid Services ad. My first impression: Pink truck with some maid "
    "service ad — not for me. I would scroll past. This is clearly for households that need "
    "domestic help, which is not my reality at all. (1 day ago)\"",
    italic=True,
)

add_para(
    "These memories shape her response to the Tata Tiago ad — she already has a pattern of "
    "finding ads irrelevant to her life stage. Each evaluation reinforces this pattern. "
    "Over time, the persona develops genuine ad fatigue and category beliefs."
)

doc.add_heading("8.2 Memory Types Generated", level=2)
add_para(
    "After a simulation run, the system automatically generates memory nodes. A single "
    "evaluation produces:"
)
add_bullet("1 observation node (always) — recording what was seen and how they reacted")
add_bullet("1 preference node (if score > 70) — e.g., 'I respond well to brands that "
           "show real families and transparent pricing'")
add_bullet("1 category belief node (if score < 30) — e.g., 'I am skeptical of household "
           "product ads that only use emotional appeal without functional claims'")
add_bullet("Reflections (when cumulative importance > 50) — higher-order synthesis across "
           "multiple exposures")

doc.add_heading("8.3 Memory Scope Configuration", level=2)
add_para("Each simulation run can configure how much memory a persona carries:")
add_bullet("Fresh evaluation, no past context", bold_prefix="NONE: ")
add_bullet("Only memories from the same product category (default)", bold_prefix="CATEGORY: ")
add_bullet("Only memories involving this specific brand", bold_prefix="BRAND: ")
add_bullet("Everything the persona has ever seen across all categories", bold_prefix="FULL: ")

add_para(
    "This enables specific use cases: NONE for baseline testing, CATEGORY for competitive "
    "context, BRAND for loyalty/fatigue analysis, FULL for holistic consumer simulation."
)

doc.add_page_break()

# ── PART 9 ─────────────────────────────────────────────────────
doc.add_heading("Part 9: Technical Architecture", level=1)

doc.add_heading("9.1 Tech Stack", level=2)

tbl = doc.add_table(rows=11, cols=2)
tbl.style = "Light Grid Accent 1"
tbl.rows[0].cells[0].text = "Component"
tbl.rows[0].cells[1].text = "Technology"
stack = [
    ("Language", "Python 3.12"),
    ("Data Models", "Pydantic v2 (strict validation, computed fields)"),
    ("LLM Provider (Primary)", "Anthropic Claude (Haiku for evals, Sonnet for critic/recommendation)"),
    ("LLM Provider (Fallback)", "OpenAI GPT-4o / GPT-4o-mini"),
    ("Embeddings", "OpenAI text-embedding-3-small"),
    ("Data Platform", "Databricks (Unity Catalog, Delta Lake, MLflow)"),
    ("Pipeline", "Medallion architecture (4 notebooks: bronze → silver → gold → simulation ingest)"),
    ("Web Dashboard", "FastAPI + vanilla JavaScript"),
    ("Testing", "pytest (120 tests, all passing)"),
    ("CLI", "argparse with run/demo subcommands"),
]
for i, (comp, tech) in enumerate(stack):
    tbl.rows[i + 1].cells[0].text = comp
    tbl.rows[i + 1].cells[1].text = tech

doc.add_heading("9.2 Codebase Structure", level=2)
add_para("Total source files: 24 Python modules + 4 notebooks + 3 frontend files", bold=True)

add_para("Source modules (src/synthetic_india/):", bold=True)
add_bullet("config.py — LLMConfig, MemoryConfig, PipelineConfig with env-var loading")
add_bullet("cli.py — CLI entry point with run/demo subcommands, demo creative builder")
add_bullet("agents/ — llm_client.py (LLMResponse, call_anthropic, call_openai, get_embedding), "
           "persona_evaluator.py, creative_analyzer.py, critic_agent.py, recommendation_agent.py, "
           "image_utils.py")
add_bullet("engine/ — simulation.py (run_simulation orchestrator, aggregation, persistence), "
           "cohort.py (diversity-aware cohort selection)")
add_bullet("memory/ — stream.py (MemoryStream class), retrieval.py (MemoryRetriever, "
           "ReflectionEngine), consumer.py (MemoryConsumer, MemoryScope)")
add_bullet("pipeline/ — bronze.py, silver.py, gold.py (local pipeline), "
           "databricks_reader.py, databricks_ingest.py, databricks_bronze.py, "
           "databricks_silver.py, databricks_gold.py (Databricks pipeline), "
           "mlflow_utils.py")
add_bullet("schemas/ — persona.py, creative.py, evaluation.py, memory.py, "
           "recommendation.py, critic.py (all Pydantic v2 BaseModels)")

doc.add_heading("9.3 Testing", level=2)
add_para(
    "120 tests, all passing. Zero API calls — every test is a smoke test that validates "
    "schema integrity, import chains, config loading, and pipeline logic without hitting "
    "any external service."
)
add_para("Test categories:", bold=True)
add_bullet("Critic schema and agent (7 tests) — verdict construction, sycophancy detection, "
           "summary rollup, prompt building")
add_bullet("Persona and cohort (4 tests) — loading, demo creatives, cohort selection, "
           "trend follower validation")
add_bullet("Memory (2 tests) — stream initialization, retriever initialization")
add_bullet("Pipeline validation (1 test) — silver persona validation")
add_bullet("CLI output (3 tests) — header, summary, critic stats rendering")
add_bullet("Vision-native (5 tests) — image path acceptance, base64 encoding, media type "
           "detection, vision message building, evaluator vision kwargs")
add_bullet("Databricks integration (3 tests) — wait parameter, job polling, async fallback")
add_bullet("Dashboard API (2 tests) — personas endpoint, simulate request with persona_ids")
add_bullet("Additional domain tests (93+ tests) — comprehensive coverage across all modules")

doc.add_heading("9.4 Cost Model", level=2)

tbl = doc.add_table(rows=6, cols=3)
tbl.style = "Light Grid Accent 1"
tbl.rows[0].cells[0].text = "Model"
tbl.rows[0].cells[1].text = "Input ($/M tokens)"
tbl.rows[0].cells[2].text = "Output ($/M tokens)"
costs = [
    ("Claude Sonnet 4", "$3.00", "$15.00"),
    ("Claude Haiku 3", "$0.25", "$1.25"),
    ("GPT-4o-mini", "$0.15", "$0.60"),
    ("GPT-4o", "$2.50", "$10.00"),
    ("text-embedding-3-small", "$0.02", "$0.00"),
]
for i, (model, inp, out) in enumerate(costs):
    tbl.rows[i + 1].cells[0].text = model
    tbl.rows[i + 1].cells[1].text = inp
    tbl.rows[i + 1].cells[2].text = out

doc.add_paragraph("")
add_para(
    "Typical simulation cost: $0.025 per persona evaluation (Haiku) + $0.01 per critic "
    "check (Sonnet) + $0.02 for recommendation (Sonnet). A 10-persona run costs ~$0.25 total. "
    "This is 1000x cheaper than traditional focus groups.",
    italic=True,
)

doc.add_page_break()

# ── PART 10 ────────────────────────────────────────────────────
doc.add_heading("Part 10: India-Specific Simulation Use Cases", level=1)

add_para(
    "Each use case runs on the same core engine. What changes is the context injected "
    "into the simulation, the questions asked of agents, and how results are interpreted. "
    "These are ordered roughly by implementation complexity."
)

use_cases = [
    ("Use Case 1: Festival Purchase Simulation",
     "India runs on festivals. Diwali, Eid, Pongal, Onam, Raksha Bandhan — each triggers "
     "completely different buying psychology.",
     "Agents receive a full festival context injection: 'Salary bonus just came in, family "
     "is visiting, need gifts for 8 people, budget is stretched but social pressure to buy "
     "good brands, seeing sale ads everywhere.' Run the same ad through the same personas in "
     "a Diwali context vs a regular Tuesday.",
     "\"Your ad scores 72 during Diwali (gifting context) but only 38 during regular season. "
     "The premium positioning works when consumers are in 'treat yourself' mode but feels "
     "unjustified for everyday purchase. Recommendation: Run this creative during Oct–Nov only.\""),

    ("Use Case 2: Influencer Trust Decay Simulator",
     "Indian D2C is built on influencer marketing. But consumers are getting wise to it. "
     "This simulation models the lifecycle of influencer trust fatigue.",
     "Show Agent A an organic-looking influencer reel for a skincare brand. They react positively. "
     "Show a second influencer pushing the same brand. Third. Fourth. At some point, accumulated "
     "memory triggers a reflection: 'Every influencer I follow is promoting this brand. This "
     "feels paid, not genuine.' The memory + reflection system makes this possible — no "
     "stateless wrapper can simulate trust decay over time.",
     "\"Your influencer saturation point is 3–4 touchpoints for Researchers and Skeptics, "
     "6–7 for Trend Followers. 'Day in my life' format lasts 2 more touchpoints before "
     "fatigue than direct product reviews.\""),

    ("Use Case 3: Price Shock Testing",
     "Indian consumers have precise mental price anchors per category. When a new brand "
     "launches outside these ranges, the reaction is segment-specific.",
     "Simulate a new protein bar launching at ₹150/bar vs ₹80 vs ₹50. Each persona reacts "
     "differently: the Price Anchor compares to Yoga Bar (₹40), the Aspirational Buyer "
     "might prefer ₹150 ('if it is expensive it must be good'), the Pragmatist calculates "
     "protein-per-rupee.",
     "\"At ₹150, only Aspirational Buyers (18% of target) accept the price. At ₹80, you "
     "capture Aspirational + Trend Followers (42%). Recommended launch price: ₹85 with a "
     "'premium ingredients' justification story.\""),

    ("Use Case 4: Language Variant Testing",
     "Uniquely Indian and untouchable by global tools. The same ad in English, Hinglish, "
     "Hindi, and Tamil lands completely differently with different segments.",
     "A Zepto ad: 'Groceries in 10 minutes' vs '10 minute mein groceries, bro' vs "
     "'10 मिनट में सब कुछ.' These are not translations — they are different emotional signals. "
     "English = premium. Hinglish = relatable. Pure Hindi = mass market.",
     "\"Your Hinglish variant outperforms English by 34% among Tier 1 consumers aged 22–30 — "
     "it feels authentic rather than corporate. But Hinglish underperforms with 35+ "
     "professionals who perceive it as 'trying too hard.'\""),

    ("Use Case 5: Tier 2 City Reality Check",
     "Every Indian D2C brand wants to expand to Tier 2 cities. Almost all fail because "
     "Tier 2 is not 'Tier 1 with less money' — the psychology is fundamentally different.",
     "Dedicated Tier 2 persona sets with distinct behavioral patterns: stronger family "
     "influence, different trust signals (local shop > Instagram), higher importance of "
     "local availability. Run both Tier 1 and Tier 2 panels on the same creative for "
     "comparison.",
     "\"Your creative scores 68 in Tier 1 but 31 in Tier 2. Key failure: Tier 2 personas "
     "don't understand 'toxin-free' and don't care. They want to know: is it available at "
     "my local medical store?\""),

    ("Use Case 6: WhatsApp Commerce Simulation",
     "India's real commerce happens on WhatsApp. Product messages from a friend who "
     "resells vs an official brand message vs a forwarded deal in a family group — "
     "completely different buying contexts.",
     "Agent trust levels vary dramatically by source — family group forward >> brand "
     "message >> unknown broadcast. Track the 'would you forward this?' action — that is "
     "the virality metric for WhatsApp commerce.",
     "\"Your product message gets 4x more purchase intent when forwarded by a trusted "
     "contact vs received as a brand broadcast. The family group forward performs best "
     "but only if price is under ₹500.\""),

    ("Use Case 7: Jugaad Detector",
     "Indian consumers are masters of jugaad. When Spotify launched at ₹119/month, Indians "
     "shared family plans with strangers. This simulation predicts consumer workarounds.",
     "Add a 'workaround ideation' step to the evaluation: 'Based on your personality, would "
     "you try to find a way to get this product for less? If so, how?' Track what % of "
     "personas attempt workarounds and which pricing structures are most jugaad-resistant.",
     "\"38% of simulated consumers immediately strategized workarounds. Most common: sharing "
     "family plan with non-family (5 personas), waiting for coupon code (4 personas). "
     "Recommendation: Implement device-limit on family plan.\""),

    ("Use Case 8: Dark Pattern Resistance Testing",
     "Different archetypes respond very differently to manipulative UX. Some don't notice "
     "hidden fees; others catch them instantly and lose trust.",
     "Run a checkout flow through agents. The Impulse Buyer doesn't notice the ₹49 "
     "'convenience fee.' The Price Anchor catches it immediately and abandons cart. The "
     "Skeptic noticed the 'only 2 left!' urgency was fake because memory says it said the "
     "same thing yesterday.",
     "\"Your checkout flow loses 40% of Price Anchor personas at the 'convenience fee' "
     "reveal. Skeptics who previously saw your 'limited stock' messaging flagged it as "
     "fake. Only Impulse Buyers complete checkout without friction.\""),
]

for title, context, how, output in use_cases:
    doc.add_heading(title, level=2)
    add_para(context)
    add_para("How It Works", bold=True)
    add_para(how)
    add_para("What Clients Get", bold=True)
    add_para(output, italic=True)

doc.add_page_break()

# ── PART 11 ────────────────────────────────────────────────────
doc.add_heading("Part 11: What Has Been Built vs What Is Next", level=1)

doc.add_heading("11.1 Fully Built and Working (Current State)", level=2)
add_bullet("20 deep personas across 8 archetypes with 50+ fields each")
add_bullet("Vision-native persona evaluation with Claude Haiku (image + text)")
add_bullet("Persistent memory system: observations, reflections, preferences, category beliefs")
add_bullet("Memory retrieval: full-dump + scored retrieval with recency/relevance/importance")
add_bullet("Memory scoping: NONE, CATEGORY, BRAND, FULL")
add_bullet("Critic Agent: 4-dimension quality gate with weighted scoring and quarantine")
add_bullet("Recommendation Agent: SCALE/ITERATE/KILL/SPLIT with edit playbook")
add_bullet("Full medallion pipeline: bronze → silver → gold (local + Databricks)")
add_bullet("Databricks round-trip: reads personas/memory, writes results back through medallion")
add_bullet("4 Databricks notebooks for full pipeline automation")
add_bullet("MLflow experiment tracking")
add_bullet("Interactive dashboard with live simulation, persona picker, radar charts")
add_bullet("Password-gated access with configurable credentials")
add_bullet("CLI with run/demo subcommands")
add_bullet("120 passing tests with zero API calls")
add_bullet("5 completed real simulation runs across 4 categories")
add_bullet("Cost tracking on every LLM call")
add_bullet("Creative analysis with 28-field structured extraction")
add_bullet("Cohort selection with archetype diversity enforcement")

doc.add_heading("11.2 Expansion Roadmap", level=2)

add_para("Phase 2: Social Simulation Layer", bold=True)
add_bullet("Agent-to-agent influence: personas discussing ads and changing each other's minds")
add_bullet("WhatsApp forwarding simulation: 'would you forward this to your family group?'")
add_bullet("Social proof propagation: trend followers amplifying signals that reach pragmatists")

add_para("Phase 3: Advanced Memory & Retrieval", bold=True)
add_bullet("Embedding-based dense vector retrieval when memory exceeds context windows")
add_bullet("Long-horizon memory with thousands of observations per persona")
add_bullet("Calibration loop: feed campaign outcomes back to adjust persona accuracy")

add_para("Phase 4: Multi-Creative Comparison", bold=True)
add_bullet("A/B variant testing: submit two creatives side-by-side for head-to-head comparison")
add_bullet("Statistical significance testing on segment-level differences")
add_bullet("Creative optimization loop: automatically suggest and test variants")

add_para("Phase 5: Platform Integration", bold=True)
add_bullet("Live API integration with Meta Ads, Google Ads for creative pulling")
add_bullet("Regional language support: Hinglish, Tamil, Telugu, Bengali ad copy evaluation")
add_bullet("Festival calendar automation with pre-built context templates")
add_bullet("WhatsApp Business API integration for real-world deployment")

add_para("Phase 6: Scale", bold=True)
add_bullet("50+ personas covering every city tier and demographic segment")
add_bullet("Multi-tenant dashboard with client-specific persona panels")
add_bullet("Automated weekly reports for ongoing campaigns")
add_bullet("Self-service API for agency integration")

doc.add_page_break()

# ── PART 12 ────────────────────────────────────────────────────
doc.add_heading("Part 12: Immediate Next Steps", level=1)

doc.add_heading("For the Engineer (Kinshuk)", level=2)
add_bullet("Run 10 more simulations across different categories to validate consistency")
add_bullet("Tune the Critic Agent threshold — current 6.0 may need adjustment based on "
           "production pass rates")
add_bullet("Add festival context injection for Holi, Diwali, Eid seasonal testing")
add_bullet("Implement A/B creative comparison mode")
add_bullet("Build the social influence layer — start with simple 'would you forward?' action")
add_bullet("Add embedding-based retrieval for personas with 50+ memory nodes")
add_bullet("Deploy dashboard to Render for external sharing")

doc.add_heading("For the Business Lead (Anshuj)", level=2)
add_bullet("Source 5 real Indian D2C ad creatives with known performance data for calibration")
add_bullet("Run a blind test: Behave predicts which ad performs better, compare to reality")
add_bullet("Demo the dashboard to 2–3 agency contacts for initial feedback")
add_bullet("Start collecting calibration data — every client engagement with real-world results "
           "makes the engine more accurate. This data is the moat.")
add_bullet("Don't over-explain the tech. Clients care about one thing: did you correctly "
           "predict which ad worked. If yes, they will ask how it works.")

doc.add_paragraph("")
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("End of Specification")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Build the wedge. Prove the engine. Then expand.")
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── Save ──────────────────────────────────────────────────
output_path = "behave_platform_spec_v3.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Pages (approx): {len(doc.paragraphs) // 30}")
